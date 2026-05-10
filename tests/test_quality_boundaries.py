import json
import logging
import sys
from types import SimpleNamespace

import pytest
from docx import Document

import ingest_service
from auth import validate_password
from chunking import chunk_document
from documents import read_docx, read_pdf_with_pdfplumber
from errors import AppError
from rag_errors import RagError
from repositories import ChatRepository
from services.document_storage import R2DocumentStorage
from services.gemini import GeminiClient
from structured_logging import JsonLogFormatter
from upload_utils import allowed_upload_extensions, safe_filename


def test_password_validation_rejects_weak_passwords():
    with pytest.raises(RagError):
        validate_password("short1")

    with pytest.raises(RagError):
        validate_password("longbutnodigits")


def test_json_log_formatter_preserves_structured_context():
    record = logging.LogRecord(
        name="tests.logger",
        level=logging.INFO,
        pathname=__file__,
        lineno=1,
        msg="user_created",
        args=(),
        exc_info=None,
    )
    record.actor_user_id = "user-1"
    record.duration_ms = 12.5

    payload = json.loads(JsonLogFormatter().format(record))

    assert payload["level"] == "info"
    assert payload["logger"] == "tests.logger"
    assert payload["message"] == "user_created"
    assert payload["actor_user_id"] == "user-1"
    assert payload["duration_ms"] == 12.5
    assert "timestamp" in payload


def test_upload_helpers_normalize_names_and_extensions(monkeypatch):
    monkeypatch.setenv("ALLOWED_UPLOAD_EXTENSIONS", "pdf, txt, docx")

    assert safe_filename("../Forest Rules?.pdf") == "Forest Rules_.pdf"
    assert allowed_upload_extensions() == {"pdf", "txt", "docx"}


def test_r2_document_storage_uses_prefixed_s3_keys(monkeypatch):
    calls = []

    class Client:
        def put_object(self, **kwargs):
            calls.append(kwargs)

    monkeypatch.setenv("R2_ACCOUNT_ID", "account-id")
    monkeypatch.setenv("R2_ACCESS_KEY_ID", "access-key")
    monkeypatch.setenv("R2_SECRET_ACCESS_KEY", "secret-key")
    monkeypatch.setenv("R2_BUCKET", "fisrag-docs")
    monkeypatch.setenv("R2_PREFIX", "source-docs")
    monkeypatch.setitem(
        sys.modules,
        "boto3",
        SimpleNamespace(client=lambda *_args, **_kwargs: Client()),
    )

    path = R2DocumentStorage().save("rules.pdf", b"content")

    assert path == "r2://fisrag-docs/source-docs/rules.pdf"
    assert calls == [{"Bucket": "fisrag-docs", "Key": "source-docs/rules.pdf", "Body": b"content"}]


def test_chunk_document_preserves_heading_context(monkeypatch):
    monkeypatch.setenv("CHUNK_TOKENS", "80")
    doc = {
        "source": "rules.txt",
        "kind": "txt",
        "title": "Forest Rules",
        "pages": [
            {
                "page": 1,
                "text": "Section 1 Introduction\n\n1. This rule applies to forest transit permits. It has clear conditions.",
            }
        ],
    }

    chunks = chunk_document(doc)

    assert chunks
    assert chunks[0]["source"] == "rules.txt"
    assert "forest transit permits" in chunks[0]["content"]


def test_read_docx_extracts_tables_as_structured_blocks(tmp_path):
    path = tmp_path / "fees.docx"
    document = Document()
    document.add_paragraph("Schedule of transit fees")
    table = document.add_table(rows=3, cols=3)
    table.rows[0].cells[0].text = "Species"
    table.rows[0].cells[1].text = "Unit"
    table.rows[0].cells[2].text = "Fee"
    table.rows[1].cells[0].text = "Teak"
    table.rows[1].cells[1].text = "Cubic meter"
    table.rows[1].cells[2].text = "1200"
    table.rows[2].cells[0].text = "Bamboo"
    table.rows[2].cells[1].text = "Bundle"
    table.rows[2].cells[2].text = "50"
    document.save(path)

    pages = read_docx(path)

    assert pages[0]["blocks"][0] == {"type": "text", "text": "Schedule of transit fees"}
    assert pages[0]["blocks"][1]["type"] == "table"
    assert pages[0]["blocks"][1]["headers"] == ["Species", "Unit", "Fee"]
    assert pages[0]["blocks"][1]["rows"][0] == ["Teak", "Cubic meter", "1200"]


def test_read_pdf_extracts_tables_as_structured_blocks(monkeypatch, tmp_path):
    class Pdf:
        pages = [
            SimpleNamespace(
                extract_text=lambda: "Schedule of transit fees",
                extract_tables=lambda: [[["Species", "Unit", "Fee"], ["Teak", "Cubic meter", "1200"]]],
            )
        ]

        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

    monkeypatch.setitem(sys.modules, "pdfplumber", SimpleNamespace(open=lambda _path: Pdf()))

    pages = read_pdf_with_pdfplumber(tmp_path / "fees.pdf")

    assert pages[0]["page"] == 1
    assert pages[0]["blocks"][0] == {"type": "text", "text": "Schedule of transit fees"}
    assert pages[0]["blocks"][1]["type"] == "table"
    assert pages[0]["blocks"][1]["headers"] == ["Species", "Unit", "Fee"]
    assert pages[0]["blocks"][1]["rows"] == [["Teak", "Cubic meter", "1200"]]


def test_chunk_document_preserves_table_header_context():
    doc = {
        "source": "fees.docx",
        "kind": "docx",
        "title": "Transit Fees",
        "pages": [
            {
                "page": None,
                "text": "Schedule 1 Transit Fees\n\nSpecies | Unit | Fee\nTeak | Cubic meter | 1200",
                "blocks": [
                    {"type": "text", "text": "Schedule 1 Transit Fees"},
                    {
                        "type": "table",
                        "table_index": 0,
                        "headers": ["Species", "Unit", "Fee"],
                        "rows": [["Teak", "Cubic meter", "1200"], ["Bamboo", "Bundle", "50"]],
                        "text": "Species | Unit | Fee\nTeak | Cubic meter | 1200\nBamboo | Bundle | 50",
                    },
                ],
            }
        ],
    }

    chunks = chunk_document(doc, max_tokens=120, overlap_tokens=0)
    table_chunks = [chunk for chunk in chunks if chunk["chunk_type"] == "table"]

    assert table_chunks
    assert "Columns: Species, Unit, Fee" in table_chunks[0]["content"]
    assert "Species: Teak" in table_chunks[0]["content"]
    assert "Unit: Cubic meter" in table_chunks[0]["content"]
    assert table_chunks[0]["metadata"]["table_indexes"] == [0]


def test_gemini_client_redacts_upstream_error_body(monkeypatch):
    class Response:
        status_code = 500
        text = "secret upstream diagnostics"

        def json(self):
            return {}

    monkeypatch.setenv("GEMINI_API_KEY", "test-key")
    monkeypatch.setattr("services.gemini.requests.post", lambda *args, **kwargs: Response())

    with pytest.raises(AppError) as exc:
        GeminiClient().generate("hello")

    assert exc.value.message == "Gemini API returned an error."
    assert exc.value.details == {"status_code": 500}
    assert "secret upstream diagnostics" not in exc.value.message


def test_chat_repository_rejects_missing_owned_session():
    class Query:
        data = []

        def select(self, *_args):
            return self

        def eq(self, *_args):
            return self

        def limit(self, *_args):
            return self

        def execute(self):
            return self

    class Client:
        def table(self, _name):
            return Query()

    with pytest.raises(RagError):
        ChatRepository(Client()).assert_session_owner("session-id", "user-id")


def test_ingest_marks_document_failed_when_chunk_insert_fails(monkeypatch):
    class Repository:
        def __init__(self):
            self.statuses = []

        def indexed_sources(self):
            return set()

        def upsert_document(self, _doc, status="indexing"):
            self.statuses.append(status)
            return "document-id"

        def replace_chunks(self, _source, _rows):
            raise RuntimeError("insert failed")

        def mark_document_status(self, _source, status, _details=None):
            self.statuses.append(status)

    repository = Repository()
    monkeypatch.setattr(
        ingest_service,
        "load_documents",
        lambda: [{"source": "a.txt", "kind": "txt", "title": "A", "page_count": None, "pages": []}],
    )
    monkeypatch.setattr(ingest_service, "chunk_document", lambda _doc: [{"source": "a.txt", "content": "x"}])
    monkeypatch.setattr(ingest_service, "chunk_row", lambda _document_id, _chunk: {"source": "a.txt"})

    with pytest.raises(RuntimeError):
        ingest_service.build_index(repository)

    assert repository.statuses == ["indexing", "failed"]
